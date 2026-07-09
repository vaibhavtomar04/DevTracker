import os
import sys
import subprocess

# Auto-install dependencies if missing
required_packages = ["python-docx", "reportlab", "markdown-it-py"]
for package in required_packages:
    try:
        __import__(package.replace("-", "_"))
    except ImportError:
        print(f"Installing missing dependency: {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

import markdown_it
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# List of doc files in order
DOC_FILES = [
    "01_executive_summary.md",
    "02_functional_spec.md",
    "03_software_spec.md",
    "04_system_architecture.md",
    "05_tech_stack.md",
    "06_project_structure.md",
    "07_database.md",
    "08_api.md",
    "09_workflows.md",
    "10_dashboards.md",
    "11_ui_components.md",
    "12_rbac.md",
    "13_security.md",
    "14_notifications.md",
    "15_file_management.md",
    "16_audit_logging.md",
    "17_reports.md",
    "18_analytics.md",
    "19_code_documentation.md",
    "20_diagrams_standards.md",
    "23_devops_testing.md",
    "26_guides_manuals.md",
    "30_roadmap.md",
]

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
MASTER_MD_PATH = os.path.join(DOCS_DIR, "DevTrack_Master_Documentation.md")
HTML_PATH = os.path.join(DOCS_DIR, "DevTrack_Documentation.html")
DOCX_PATH = os.path.join(DOCS_DIR, "DevTrack_Documentation.docx")
PDF_PATH = os.path.join(DOCS_DIR, "DevTrack_Documentation.pdf")

def build_master_markdown():
    print("Concatenating files into Master Markdown...")
    master_content = []
    master_content.append("# DevTrack 2.0 Enterprise Technical Reference Suite\n")
    master_content.append("## Complete Documentation and System Specifications\n")
    master_content.append("---\n")
    
    for filename in DOC_FILES:
        filepath = os.path.join(DOCS_DIR, filename)
        if os.path.exists(filepath):
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
                master_content.append(content)
                master_content.append("\n\n---\n\n")
        else:
            print(f"Warning: {filename} not found.")
            
    with open(MASTER_MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(master_content))
    print(f"[SUCCESS] Master Markdown saved to: {MASTER_MD_PATH}")
    return "\n".join(master_content)

def build_html(master_content):
    print("Building HTML Documentation...")
    md = markdown_it.MarkdownIt()
    html_body = md.render(master_content)
    
    # Premium glassmorphic/dark theme CSS with clean printing styles
    css_styles = """
    :root {
        --bg-color: #060814;
        --card-bg: rgba(255, 255, 255, 0.03);
        --border-color: rgba(255, 255, 255, 0.08);
        --text-primary: #f1f5f9;
        --text-secondary: #94a3b8;
        --accent-violet: #8b5cf6;
        --accent-cyan: #22d3ee;
        --code-bg: rgba(255, 255, 255, 0.05);
    }
    body {
        font-family: 'Inter', system-ui, sans-serif;
        background-color: var(--bg-color);
        color: var(--text-primary);
        line-height: 1.7;
        margin: 0;
        padding: 40px 20px;
        display: flex;
        justify-content: center;
    }
    .container {
        max-width: 900px;
        width: 100%;
        background-color: var(--card-bg);
        border: 1px solid var(--border-color);
        border-radius: 20px;
        padding: 50px;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.5);
        backdrop-filter: blur(8px);
    }
    h1, h2, h3, h4 {
        color: var(--text-primary);
        font-weight: 800;
        margin-top: 1.5em;
        letter-spacing: -0.02em;
    }
    h1 {
        font-size: 2.2em;
        border-bottom: 2px solid var(--accent-violet);
        padding-bottom: 10px;
        margin-bottom: 1em;
        background: linear-gradient(to right, #a78bfa, #22d3ee);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    h2 {
        font-size: 1.6em;
        color: #c084fc;
        border-bottom: 1px solid var(--border-color);
        padding-bottom: 5px;
    }
    h3 {
        font-size: 1.2em;
        color: var(--accent-cyan);
    }
    p, li {
        color: var(--text-secondary);
        font-size: 14px;
    }
    a {
        color: var(--accent-cyan);
        text-decoration: none;
    }
    a:hover {
        text-decoration: underline;
    }
    pre {
        background-color: var(--code-bg);
        border: 1px solid var(--border-color);
        padding: 15px;
        border-radius: 10px;
        overflow-x: auto;
    }
    code {
        font-family: 'Fira Code', Consolas, monospace;
        font-size: 13px;
        color: #f472b6;
    }
    table {
        width: 100%;
        border-collapse: collapse;
        margin: 20px 0;
        font-size: 13px;
    }
    th, td {
        padding: 12px;
        border: 1px solid var(--border-color);
        text-align: left;
    }
    th {
        background-color: rgba(139, 92, 246, 0.1);
        color: #c084fc;
    }
    td {
        color: var(--text-secondary);
    }
    @media print {
        body {
            background-color: #fff;
            color: #000;
            padding: 0;
        }
        .container {
            border: none;
            box-shadow: none;
            padding: 0;
            max-width: 100%;
        }
        h1, h2, h3, h4, p, li, td, th {
            color: #000 !important;
            background: none !important;
            -webkit-text-fill-color: initial !important;
        }
        pre {
            background-color: #f5f5f5;
            border: 1px solid #ccc;
            page-break-inside: avoid;
        }
        table {
            page-break-inside: avoid;
        }
    }
    """
    
    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>DevTrack 2.0 Enterprise Technical Reference</title>
    <style>{css_styles}</style>
</head>
<body>
    <div class="container">
        {html_body}
    </div>
</body>
</html>
"""
    with open(HTML_PATH, "w", encoding="utf-8") as f:
        f.write(full_html)
    print(f"[SUCCESS] HTML Document saved to: {HTML_PATH}")

def build_docx():
    print("Building Word DOCX Document...")
    doc = Document()
    
    # Configure styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("DEVTRACK 2.0\nENTERPRISE TECHNICAL SUITE")
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x7c, 0x3a, 0xed) # Violet
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("System Documentation and Technical Reference Manual")
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    
    doc.add_page_break()
    
    # Parse Master Markdown to populate DOCX
    with open(MASTER_MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            run = h.add_run(stripped[2:])
            run.font.color.rgb = RGBColor(0x7c, 0x3a, 0xed)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            run = h.add_run(stripped[3:])
            run.font.color.rgb = RGBColor(0x0e, 0x74, 0x90) # Cyan-blue
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            run = h.add_run(stripped[4:])
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(stripped[2:])
        else:
            doc.add_paragraph(stripped)
            
    doc.save(DOCX_PATH)
    print(f"[SUCCESS] DOCX Document saved to: {DOCX_PATH}")

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        if self._pageNumber == 1:
            return  # Suppress page number on cover page
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#64748b"))
        self.drawString(54, 30, "DevTrack 2.0 Technical Reference Manual")
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(612 - 54, 30, page_str)
        self.setStrokeColor(colors.HexColor("#e2e8f0"))
        self.setLineWidth(0.5)
        self.line(54, 45, 612 - 54, 45)
        self.restoreState()

def build_pdf():
    print("Building PDF Document...")
    doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Styles
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=30,
        textColor=colors.HexColor("#7c3aed"),
        alignment=1, # Center
        spaceAfter=15
    )
    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#64748b"),
        alignment=1,
        spaceAfter=150
    )
    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#7c3aed"),
        spaceBefore=15,
        spaceAfter=10,
        keepWithNext=True
    )
    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#0f766e"),
        spaceBefore=12,
        spaceAfter=8,
        keepWithNext=True
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#334155"),
        spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#334155"),
        leftIndent=20,
        spaceAfter=4
    )

    story = []
    
    # Title Page
    story.append(Spacer(1, 100))
    story.append(Paragraph("DEVTRACK 2.0", title_style))
    story.append(Paragraph("ENTERPRISE TECHNICAL DOCUMENTATION", title_style))
    story.append(Paragraph("A Comprehensive Guide to Architecture, Workflows, APIs, and Specifications", subtitle_style))
    story.append(Paragraph("Version 1.0.0 &bull; Confidential &bull; Standard Technical Reference", ParagraphStyle('CenterMuted', parent=styles['Normal'], alignment=1, textColor=colors.HexColor("#94a3b8"))))
    story.append(PageBreak())
    
    # Parse Master MD to build PDF elements
    with open(MASTER_MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    import html
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            story.append(Paragraph(html.escape(stripped[2:]), h1_style))
        elif stripped.startswith("## "):
            story.append(Paragraph(html.escape(stripped[3:]), h2_style))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            story.append(Paragraph(f"&bull; {html.escape(stripped[2:])}", bullet_style))
        else:
            story.append(Paragraph(html.escape(stripped), body_style))
            
    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"[SUCCESS] PDF Document saved to: {PDF_PATH}")

if __name__ == "__main__":
    print("Starting documentation compilation script...")
    content = build_master_markdown()
    build_html(content)
    build_docx()
    build_pdf()
    print("==================================================")
    print("DevTrack 2.0 Documentation Suite successfully built!")
    print("==================================================")
